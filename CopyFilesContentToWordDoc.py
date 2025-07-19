import os
from docx import Document
import logging

# === CONFIGURATION ===
root_dir = r"D:\AlphaNivesh\ANQuant"
allowed_extensions = {'.py', '.yaml', '.yml', '.rs', '.toml'}
output_file = "project_files_dump.docx"
exclude_dirs = {'.git', '.venv', '__pycache__', 'site-packages', 'build', 'dist', 'node_modules'}

# === SETUP LOGGING ===
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s"
)

logging.info("🚀 Starting file extraction and Word document creation...")
logging.info(f"📁 Scanning root directory: {root_dir}")

# === CREATE DOCUMENT ===
doc = Document()
doc.add_heading("Project Files Dump", level=0)

file_count = 0
error_count = 0

# === WALK THROUGH FILES ===
for dirpath, dirnames, filenames in os.walk(root_dir):
    # Modify dirnames in-place to skip excluded folders
    dirnames[:] = [d for d in dirnames if d not in exclude_dirs and not d.startswith('.')]

    for file in filenames:
        _, ext = os.path.splitext(file)
        if ext.lower() in allowed_extensions:
            full_path = os.path.join(dirpath, file)
            rel_path = os.path.relpath(full_path, root_dir)

            logging.info(f"📄 Processing file: {rel_path}")
            doc.add_heading(rel_path, level=2)

            try:
                with open(full_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                doc.add_paragraph(content, style='Normal')
                file_count += 1
            except Exception as e:
                logging.error(f"❌ Failed to read file {rel_path}: {e}")
                doc.add_paragraph(f"[Error reading file: {e}]")
                error_count += 1

# === SAVE DOC ===
try:
    doc.save(output_file)
    logging.info(f"✅ Word document saved successfully as: {output_file}")
    logging.info(f"📊 Total files processed: {file_count}, Errors: {error_count}")
except Exception as e:
    logging.error(f"❌ Failed to save the Word document: {e}")
